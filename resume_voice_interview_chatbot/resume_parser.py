"""
Resume Parser Module
====================
Parses PDF and DOCX resume files to extract structured information
including personal details, education, skills, experience, and more.
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field, asdict
from loguru import logger

# Import parsing libraries
import fitz  # PyMuPDF for PDF
from docx import Document


@dataclass
class PersonalInfo:
    """Personal information extracted from resume."""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None
    website: Optional[str] = None


@dataclass
class Education:
    """Education entry."""
    degree: Optional[str] = None
    institution: Optional[str] = None
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    gpa: Optional[str] = None
    description: Optional[str] = None


@dataclass
class Experience:
    """Work experience entry."""
    title: Optional[str] = None
    company: Optional[str] = None
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    current: bool = False
    description: Optional[str] = None
    achievements: List[str] = field(default_factory=list)


@dataclass
class Project:
    """Project entry."""
    name: Optional[str] = None
    description: Optional[str] = None
    technologies: List[str] = field(default_factory=list)
    url: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    highlights: List[str] = field(default_factory=list)


@dataclass
class Certification:
    """Certification entry."""
    name: Optional[str] = None
    issuer: Optional[str] = None
    date: Optional[str] = None
    expiry: Optional[str] = None
    credential_id: Optional[str] = None
    url: Optional[str] = None


@dataclass
class Skill:
    """Skill entry."""
    name: str
    level: Optional[str] = None
    category: Optional[str] = None


@dataclass
class Achievement:
    """Achievement entry."""
    title: Optional[str] = None
    description: Optional[str] = None
    date: Optional[str] = None
    issuer: Optional[str] = None


class ResumeParser:
    """Main resume parser class."""
    
    # Regex patterns for extraction
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'(\+?91[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    LINKEDIN_PATTERN = r'linkedin\.com/in/[a-zA-Z0-9_-]+'
    GITHUB_PATTERN = r'github\.com/[a-zA-Z0-9_-]+'
    
    # Section headers
    SECTION_HEADERS = {
        'education': ['education', 'academic', 'qualification', 'degree'],
        'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience'],
        'projects': ['projects', 'personal projects', 'academic projects', 'side projects', 
                     'project experience', 'academic project', 'portfolio', 'work samples',
                     'research projects', 'university projects', 'internship projects'],
        'skills': ['skills', 'technical skills', 'technologies', 'competencies', 'expertise'],
        'certifications': ['certifications', 'certificates', 'credentials'],
        'achievements': ['achievements', 'awards', 'recognitions', 'accomplishments'],
        'summary': ['summary', 'objective', 'profile', 'about'],
    }
    
    # Skill categories
    SKILL_CATEGORIES = {
        'programming_languages': [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 
            'rust', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'shell',
            'bash', 'powershell', 'sql', 'html', 'css', 'sass', 'less'
        ],
        'frameworks': [
            'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'spring', 'node.js',
            'express', 'rails', 'laravel', 'flutter', 'react native', 'next.js', 'nuxt',
            '.net', 'asp.net', 'hibernate', 'spring boot', 'tensorflow', 'pytorch',
            'keras', 'scikit-learn', 'pandas', 'numpy'
        ],
        'databases': [
            'mysql', 'postgresql', 'mongodb', 'redis', 'cassandra', 'elasticsearch',
            'oracle', 'sql server', 'sqlite', 'dynamodb', 'firebase', 'supabase'
        ],
        'cloud': [
            'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'docker',
            'kubernetes', 'terraform', 'ansible', 'jenkins', 'github actions', 'gitlab ci'
        ],
        'ml_ai': [
            'machine learning', 'deep learning', 'nlp', 'computer vision', 'tensorflow',
            'pytorch', 'keras', 'scikit-learn', 'opencv', 'spacy', 'nltk', 'hugging face',
            'transformers', 'llm', 'gpt', 'bert', 'reinforcement learning', 'xgboost'
        ],
        'tools': [
            'git', 'github', 'gitlab', 'jira', 'confluence', 'slack', 'postman',
            'tableau', 'powerbi', 'excel', 'figma', 'sketch', 'adobe xd'
        ]
    }
    
    def __init__(self):
        """Initialize resume parser."""
        self.raw_text = ""
        self.lines = []
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(str(file_path))
            text_parts = []
            
            for page in doc:
                text_parts.append(page.get_text())
            
            doc.close()
            
            logger.info(f"Extracted text from PDF: {file_path.name}")
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            logger.info(f"Extracted text from DOCX: {file_path.name}")
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from resume file (auto-detect type)."""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse resume and extract all information.
        
        Returns:
            Dictionary containing extracted resume data
        """
        logger.info(f"Parsing resume: {file_path.name}")
        
        # Extract raw text
        self.raw_text = self.extract_text(file_path)
        self.lines = [line.strip() for line in self.raw_text.split('\n') if line.strip()]
        
        # Extract all sections
        parsed_data = {
            'personal_info': self._extract_personal_info(),
            'summary': self._extract_summary(),
            'education': self._extract_education(),
            'experience': self._extract_experience(),
            'projects': self._extract_projects(),
            'skills': self._extract_skills(),
            'certifications': self._extract_certifications(),
            'achievements': self._extract_achievements(),
        }
        
        # Add extracted skills list
        parsed_data['extracted_skills'] = self._get_all_skills(parsed_data)
        
        logger.info(f"Resume parsed successfully: {file_path.name}")
        return parsed_data
    
    def _extract_personal_info(self) -> Dict[str, Any]:
        """Extract personal information."""
        personal_info = PersonalInfo()
        
        # Extract email
        email_match = re.search(self.EMAIL_PATTERN, self.raw_text)
        if email_match:
            personal_info.email = email_match.group()
        
        # Extract phone
        phone_match = re.search(self.PHONE_PATTERN, self.raw_text)
        if phone_match:
            personal_info.phone = phone_match.group()
        
        # Extract LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, self.raw_text, re.IGNORECASE)
        if linkedin_match:
            personal_info.linkedin = f"https://{linkedin_match.group()}"
        
        # Extract GitHub
        github_match = re.search(self.GITHUB_PATTERN, self.raw_text, re.IGNORECASE)
        if github_match:
            personal_info.github = f"https://{github_match.group()}"
        
        # First non-empty line is often the name
        for line in self.lines[:5]:
            # Skip if it looks like a header or common word
            if any(header in line.lower() for header in 
                   ['resume', 'cv', 'curriculum', 'contact', 'objective', 'summary']):
                continue
            # Skip if it's just contact info
            if re.search(self.EMAIL_PATTERN, line) or re.search(self.PHONE_PATTERN, line):
                continue
            # Check if it looks like a name (2-4 words, mostly letters)
            words = line.split()
            if 1 < len(words) <= 4:
                if all(re.match(r'^[A-Za-z\s\'-]+$', word) for word in words):
                    personal_info.name = line
                    break
        
        # Try to extract location
        location_patterns = [
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*([A-Z]{2})',  # City, ST
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)?),\s*([A-Z][a-z]+)',  # City, Country
            r'[📍📎]\s*([A-Za-z\s,]+)',
        ]
        
        for pattern in location_patterns:
            location_match = re.search(pattern, self.raw_text)
            if location_match:
                personal_info.location = location_match.group(1)
                break
        
        return asdict(personal_info)
    
    def _extract_summary(self) -> Optional[str]:
        """Extract professional summary."""
        return self._extract_section('summary')
    
    def _extract_section(self, section_type: str) -> Optional[str]:
        """Generic section extraction."""
        headers = self.SECTION_HEADERS.get(section_type, [])
        
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                # Get content until next section header
                content_lines = []
                for j in range(i + 1, min(i + 10, len(self.lines))):
                    next_line = self.lines[j]
                    # Check if it's a section header
                    is_header = False
                    for headers_list in self.SECTION_HEADERS.values():
                        if any(h in next_line.lower() for h in headers_list):
                            is_header = True
                            break
                    if is_header:
                        break
                    content_lines.append(next_line)
                
                if content_lines:
                    return ' '.join(content_lines)
        
        return None
    
    def _extract_education(self) -> List[Dict[str, Any]]:
        """Extract education entries."""
        education_list = []
        headers = self.SECTION_HEADERS['education']
        
        in_section = False
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line.lower() for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    break
                
                # Try to parse education entry
                edu = self._parse_education_entry(line)
                if edu:
                    education_list.append(edu)
        
        return education_list
    
    def _parse_education_entry(self, text: str) -> Optional[Dict[str, Any]]:
        """Parse a single education entry."""
        education = Education()
        
        # Common degree patterns
        degree_patterns = [
            r'(B\.?S\.?|B\.?A\.?|Bachelor)',
            r'(M\.?S\.?|M\.?A\.?|Master)',
            r'(Ph\.?D\.?|Doctor)',
            r'(B\.?Tech|M\.?Tech)',
            r'(MBA)',
        ]
        
        for pattern in degree_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                education.degree = match.group()
                break
        
        # Date patterns
        date_pattern = r'(\d{4})\s*[-–]\s*(\d{4}|\w+)'
        date_match = re.search(date_pattern, text)
        if date_match:
            education.start_date = date_match.group(1)
            education.end_date = date_match.group(2)
        
        # GPA pattern
        gpa_pattern = r'GPA[:\s]*(\d+\.?\d*)'
        gpa_match = re.search(gpa_pattern, text, re.IGNORECASE)
        if gpa_match:
            education.gpa = gpa_match.group(1)
        
        if education.degree:
            return asdict(education)
        return None
    
    def _extract_experience(self) -> List[Dict[str, Any]]:
        """Extract work experience entries."""
        experience_list = []
        headers = self.SECTION_HEADERS['experience']
        
        in_section = False
        current_entry = None
        entry_lines = []
        
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line.lower() for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    if current_entry:
                        parsed = self._parse_experience_entry(entry_lines)
                        if parsed:
                            experience_list.append(parsed)
                    break
                
                # Check if line looks like a new job entry
                # (has company name patterns, dates, or titles)
                if self._looks_like_job_header(line):
                    if current_entry:
                        parsed = self._parse_experience_entry(entry_lines)
                        if parsed:
                            experience_list.append(parsed)
                    current_entry = line
                    entry_lines = [line]
                elif current_entry:
                    entry_lines.append(line)
        
        # Don't forget the last entry
        if entry_lines:
            parsed = self._parse_experience_entry(entry_lines)
            if parsed:
                experience_list.append(parsed)
        
        return experience_list
    
    def _looks_like_job_header(self, line: str) -> bool:
        """Check if a line looks like a job header."""
        # Has date range
        if re.search(r'\d{4}\s*[-–]\s*(Present|\d{4})', line):
            return True
        # Has common job title keywords
        title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer', 
                        'consultant', 'architect', 'lead', 'senior', 'junior', 'intern',
                        'associate', 'director', 'head', 'chief']
        if any(kw in line.lower() for kw in title_keywords):
            return True
        return False
    
    def _parse_experience_entry(self, lines: List[str]) -> Optional[Dict[str, Any]]:
        """Parse a single experience entry."""
        experience = Experience()
        full_text = ' '.join(lines)
        
        # Extract job title
        title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer', 
                        'consultant', 'architect', 'lead', 'senior', 'junior', 'intern',
                        'associate', 'director', 'head', 'chief']
        
        for line in lines:
            for kw in title_keywords:
                if kw in line.lower():
                    # Find the full title
                    words = line.split()
                    for i, word in enumerate(words):
                        if kw in word.lower():
                            # Get surrounding context
                            start = max(0, i - 2)
                            end = min(len(words), i + 4)
                            experience.title = ' '.join(words[start:end])
                            break
                    break
            if experience.title:
                break
        
        # Extract dates
        date_pattern = r'(\w+\s*\d{4})\s*[-–]\s*(\w+\s*\d{4}|Present)'
        date_match = re.search(date_pattern, full_text)
        if date_match:
            experience.start_date = date_match.group(1)
            end_date = date_match.group(2)
            experience.current = 'present' in end_date.lower()
            experience.end_date = end_date
        
        # Extract company (try to find after @ or | or in parens)
        company_patterns = [
            r'@\s*([A-Z][a-zA-Z\s]+)',
            r'\|\s*([A-Z][a-zA-Z\s]+)',
            r'\(([A-Z][a-zA-Z\s]+)\)',
        ]
        
        for pattern in company_patterns:
            match = re.search(pattern, full_text)
            if match:
                experience.company = match.group(1).strip()
                break
        
        # Extract achievements (bullet points)
        achievement_keywords = ['increased', 'reduced', 'led', 'achieved', 'improved',
                              'implemented', 'developed', 'created', 'managed', 'won',
                              'delivered', 'launched', 'saved', 'generated']
        
        achievements = []
        for line in lines:
            if any(kw in line.lower() for kw in achievement_keywords):
                # Clean the line
                clean_line = re.sub(r'^[\s•\-\*\d\.]+', '', line).strip()
                if clean_line:
                    achievements.append(clean_line)
        
        experience.achievements = achievements
        experience.description = full_text
        
        if experience.title:
            return asdict(experience)
        return None
    
    def _extract_projects(self) -> List[Dict[str, Any]]:
        """Extract project entries."""
        projects_list = []
        headers = self.SECTION_HEADERS['projects']
        
        in_section = False
        project_lines = []
        
        for i, line in enumerate(self.lines):
            line_lower = line.lower().strip()
            
            # Check for project headers
            if any(header in line_lower for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line_lower for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    if project_lines:
                        parsed = self._parse_project_entry(project_lines)
                        if parsed:
                            projects_list.append(parsed)
                    break
                
                # Check if line looks like a new project
                # Project names often have: | (pipe), : (colon), start with capital letter in a line, or are numbered
                is_new_project = False
                stripped = line.strip()
                
                if '|' in stripped or ':' in stripped:
                    is_new_project = True
                elif stripped and stripped[0].isupper() and len(stripped) < 80:
                    # Likely a title/name of a project
                    is_new_project = True
                elif stripped and (stripped[0].isdigit() or stripped.startswith('- ') or stripped.startswith('* ')):
                    is_new_project = True
                
                if is_new_project:
                    if project_lines:
                        parsed = self._parse_project_entry(project_lines)
                        if parsed:
                            projects_list.append(parsed)
                    project_lines = [line]
                elif project_lines:
                    project_lines.append(line)
        
        # Don't forget the last project
        if project_lines:
            parsed = self._parse_project_entry(project_lines)
            if parsed:
                projects_list.append(parsed)
        
        return projects_list
    
    def _parse_project_entry(self, lines: List[str]) -> Optional[Dict[str, Any]]:
        """Parse a single project entry."""
        project = Project()
        full_text = ' '.join(lines)
        
        # Extract project name (usually first significant line or before | or :)
        first_line = lines[0] if lines else ""
        if '|' in first_line:
            project.name = first_line.split('|')[0].strip()
        elif ':' in first_line:
            project.name = first_line.split(':')[0].strip()
        else:
            project.name = first_line.strip()
        
        # Extract technologies
        tech_keywords = []
        for category, skills in self.SKILL_CATEGORIES.items():
            for skill in skills:
                if skill.lower() in full_text.lower():
                    tech_keywords.append(skill)
        project.technologies = list(set(tech_keywords))
        
        # Extract description
        descriptions = []
        for line in lines[1:]:
            clean_line = re.sub(r'^[\s•\-\*\d\.]+', '', line).strip()
            if clean_line:
                descriptions.append(clean_line)
        project.description = ' '.join(descriptions) if descriptions else None
        
        # Extract dates
        date_pattern = r'(\w+\s*\d{4})\s*[-–]\s*(\w+\s*\d{4}|Present)'
        date_match = re.search(date_pattern, full_text)
        if date_match:
            project.start_date = date_match.group(1)
            project.end_date = date_match.group(2)
        
        if project.name:
            return asdict(project)
        return None
    
    def _extract_skills(self) -> List[Dict[str, Any]]:
        """Extract skills with categories."""
        skills_list = []
        headers = self.SECTION_HEADERS['skills']
        
        in_section = False
        skill_lines = []
        
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line.lower() for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    break
                
                skill_lines.append(line)
        
        # Parse skills from collected lines
        for line in skill_lines:
            # Split by common delimiters
            parts = re.split(r'[,;•|\n]', line)
            for part in parts:
                skill_name = part.strip()
                if skill_name and len(skill_name) > 1:
                    # Determine category
                    category = self._categorize_skill(skill_name)
                    skill = Skill(name=skill_name, category=category)
                    skills_list.append(asdict(skill))
        
        return skills_list
    
    def _categorize_skill(self, skill: str) -> Optional[str]:
        """Categorize a skill."""
        skill_lower = skill.lower()
        
        for category, skills in self.SKILL_CATEGORIES.items():
            for s in skills:
                if s in skill_lower or skill_lower in s:
                    return category
        
        return 'other'
    
    def _extract_certifications(self) -> List[Dict[str, Any]]:
        """Extract certifications."""
        certifications_list = []
        headers = self.SECTION_HEADERS['certifications']
        
        in_section = False
        
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line.lower() for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    break
                
                cert = self._parse_certification(line)
                if cert:
                    certifications_list.append(cert)
        
        return certifications_list
    
    def _parse_certification(self, text: str) -> Optional[Dict[str, Any]]:
        """Parse a single certification entry."""
        cert = Certification()
        
        # Extract certification name (before common separators)
        if '|' in text:
            parts = text.split('|')
            cert.name = parts[0].strip()
            if len(parts) > 1:
                cert.issuer = parts[1].strip()
        elif '-' in text:
            parts = text.split('-')
            cert.name = parts[0].strip()
            if len(parts) > 1:
                cert.issuer = parts[1].strip()
        else:
            cert.name = text.strip()
        
        # Extract date
        date_pattern = r'(\w+\s*\d{4})'
        date_match = re.search(date_pattern, text)
        if date_match:
            cert.date = date_match.group(1)
        
        if cert.name:
            return asdict(cert)
        return None
    
    def _extract_achievements(self) -> List[Dict[str, Any]]:
        """Extract achievements and awards."""
        achievements_list = []
        headers = self.SECTION_HEADERS['achievements']
        
        in_section = False
        
        for i, line in enumerate(self.lines):
            if any(header in line.lower() for header in headers):
                in_section = True
                continue
            
            if in_section:
                # Check if we've hit another section
                is_section = False
                for headers_list in self.SECTION_HEADERS.values():
                    if any(h in line.lower() for h in headers_list):
                        is_section = True
                        break
                
                if is_section:
                    break
                
                achievement = Achievement()
                achievement.description = re.sub(r'^[\s•\-\*\d\.]+', '', line).strip()
                
                # Extract date
                date_pattern = r'(\w+\s*\d{4})'
                date_match = re.search(date_pattern, line)
                if date_match:
                    achievement.date = date_match.group(1)
                
                if achievement.description:
                    achievements_list.append(asdict(achievement))
        
        return achievements_list
    
    def _get_all_skills(self, parsed_data: Dict) -> List[str]:
        """Get flat list of all extracted skills."""
        skills = set()
        
        # From skills section
        for skill in parsed_data.get('skills', []):
            if skill.get('name'):
                skills.add(skill['name'])
        
        # From experience
        for exp in parsed_data.get('experience', []):
            desc = exp.get('description', '')
            for category, skill_list in self.SKILL_CATEGORIES.items():
                for skill in skill_list:
                    if skill.lower() in desc.lower():
                        skills.add(skill)
        
        # From projects
        for proj in parsed_data.get('projects', []):
            for tech in proj.get('technologies', []):
                skills.add(tech)
        
        return sorted(list(skills))


def parse_resume(file_path: Path) -> Dict[str, Any]:
    """Convenience function to parse a resume file."""
    parser = ResumeParser()
    return parser.parse(file_path)


# Example usage
if __name__ == "__main__":
    # Test parsing
    import sys
    
    if len(sys.argv) > 1:
        file_path = Path(sys.argv[1])
        parser = ResumeParser()
        result = parser.parse(file_path)
        print(json.dumps(result, indent=2))
